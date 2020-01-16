import docx
from datetime import datetime
import pandas as pd
from pathlib import Path
import os
import yaml

time_start = datetime.now()

project_config = docx.Document("config.docx")

weekly_source_path = Path(project_config.paragraphs[1].text.split("=")[1].replace("\"", ""))
output_path = Path(project_config.paragraphs[2].text.split("=")[1].replace("\"", ""))
banks_names = yaml.safe_load(project_config.paragraphs[5].text.split("=")[1].replace("\"", ""))

codes_to_remove = ["АВТ.КОД/АС:", "Авт. код: ", "Авт. код ", "Авт.код :", "Авт. код. ", "АВТ. КОД/ ", "АВТ.КОД/AC:",
                   "AUTH.CODE:", "AUTH CODE:", "Auth. Code ", "PRE-AUTH ", "ACC", "AC:", "АС:", "AC :", "АС :", "AC: ",
                   "AC", "АС", "AС", "АC", "AC ", "ac ", "Ас ", "ac", "Ac", "tid ", "Авт.код:", ]

os.chdir(output_path)

week_to_process = int(input("Week to process: "))
winners_to_draw = int(input("Winners: "))
reserves_to_draw = int(input("Reserves: "))

entries_to_sample = winners_to_draw + reserves_to_draw

for weekly_folder_element in weekly_source_path.glob("*"):
    if os.path.isdir(weekly_folder_element) and weekly_folder_element.name != "_Results" and int(weekly_folder_element.name.split(".")[0]) == week_to_process:
        for file in weekly_folder_element.iterdir():
            if file.suffix == ".xlsx":
                all_data_frame = pd.read_excel(file, sheet_name=0, skiprows=None)
                all_data_frame.columns = all_data_frame.columns.str.strip()
                all_data_frame["Име*:"] = all_data_frame["Име*:"].str.strip()

                for code in codes_to_remove:
                    all_data_frame["Отор. код на ПОС бележка*:"] = [x.strip().replace(code, '').strip() for x in all_data_frame["Отор. код на ПОС бележка*:"]]

                duplicates_data_frame = all_data_frame[
                    all_data_frame.duplicated(["Отор. код на ПОС бележка*:", "Имейл*:"],
                                              keep="first")]
                all_data_frame.drop_duplicates(["Отор. код на ПОС бележка*:", "Имейл*:"],
                                               keep="first", inplace=True)

                all_drawn_data_frame = all_data_frame.sample(n=entries_to_sample)

                whole_week = weekly_folder_element.name.split()[1]

                if week_to_process < 10:
                    week_to_process = "0" + str(week_to_process)

                if not os.path.exists(f"Week_{week_to_process}"):
                    os.makedirs(f"Week_{week_to_process}")
                if not os.path.exists(f"Week_{week_to_process}/Banks"):
                    os.makedirs(f"Week_{week_to_process}/Banks")

                weekly_winners = f"Week_{week_to_process}/Week_{whole_week}_winners.xlsx"
                weekly_reserves = f"Week_{week_to_process}/Week_{whole_week}_reserves.xlsx"
                weekly_duplicates = f"Week_{week_to_process}/Week_{whole_week}_duplicates.xlsx"
                weekly_no_duplicates = f"Week_{week_to_process}/Week_{whole_week}_no_duplicates.xlsx"

                if os.path.exists(weekly_winners):
                    os.remove(weekly_winners)
                if os.path.exists(weekly_reserves):
                    os.remove(weekly_reserves)
                if os.path.exists(weekly_duplicates):
                    os.remove(weekly_duplicates)
                if os.path.exists(weekly_no_duplicates):
                    os.remove(weekly_no_duplicates)

                all_drawn_data_frame = all_drawn_data_frame.reset_index(drop=True)

                winners_data_frame = all_drawn_data_frame.loc[0:(winners_to_draw - 1)]
                reserves_data_frame = all_drawn_data_frame.loc[winners_to_draw:(winners_to_draw + reserves_to_draw - 1)]

                winners_data_frame.to_excel(weekly_winners, index=False)
                reserves_data_frame.to_excel(weekly_reserves, index=False)
                duplicates_data_frame.to_excel(weekly_duplicates, index=False)
                all_data_frame.to_excel(weekly_no_duplicates, index=False)

                os.chdir(Path(os.path.join(os.getcwd(), f"Week_{week_to_process}", "Banks")))

                banks_list = winners_data_frame["Банка*:"].unique()

                for bank in banks_list:
                    win_bank_data_frame = winners_data_frame[winners_data_frame["Банка*:"] == bank]
                    bank_name = banks_names[bank]
                    if win_bank_data_frame.shape[0] == 1:
                        win_bank_file = f"Week_{week_to_process}_winner_{bank_name}.xlsx"
                    else:
                        win_bank_file = f"Week_{week_to_process}_winners_{bank_name}.xlsx"
                    if os.path.isfile(win_bank_file):
                        os.remove(win_bank_file)
                    win_bank_data_frame.to_excel(win_bank_file, index=False)

time_end = datetime.now()
time_took = time_end - time_start

print(f"The execution of this script took {time_took.seconds} seconds.")


