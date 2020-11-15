from datetime import datetime
from pathlib import Path
import docx
import pandas as pd
import os
import yaml


def append_and_print_statistics(data, li):
    print(data)
    li.append(data)
    return li


def process_code(data, current_code):
    symbols = ["/", "NO"]
    data = data.strip()
    if current_code in data[0: len(current_code)]:
        data = data.strip().replace(current_code, "").strip()
        for symbol in symbols:
            if symbol in data:
                data = data.split(symbol)[0].strip()

    return data


time_start = datetime.now()

project_config = docx.Document("config_new.docx")

source_path = Path(project_config.paragraphs[0].text.split("=")[1].replace("\"", ""))  # pathlib.WindowsPath
output_path = Path(project_config.paragraphs[1].text.split("=")[1].replace("\"", ""))   # pathlib.WindowsPath
banks_en_names = yaml.safe_load(project_config.paragraphs[2].text.split("=")[1].replace("\"", ""))  # dict
codes_to_remove = project_config.paragraphs[3].text.split("=")[1].strip().replace("\"", "").split(", ")  # list

statistics = []

week_to_process = 2
weekly_winners = 50
weekly_reserves = 50

whole_week = ""

os.chdir(output_path)

for element in source_path.glob("*"):
    if os.path.isdir(element):
        if element.name != "_Results" and int(element.name.split(".")[0]) == week_to_process:
            for file in element.iterdir():
                if file.suffix == ".xlsx":

                    all_df = pd.read_excel(file, sheet_name=0, skiprows=None)
                    all_df.columns = all_df.columns.str.strip()
                    all_df["Банка*:"] = all_df["Банка*:"].str.strip()
                    all_df["Име*:"] = all_df["Име*:"].str.strip()
                    all_df[["Reg_date", "Reg_time"]] = all_df["Submitted date"].str.split(" ", expand=True)

                    for code in codes_to_remove:
                        all_df["Отор. код на ПОС бележка*:"] = \
                            [process_code(x, code) for x in all_df["Отор. код на ПОС бележка*:"]]

                    initial_df = all_df.copy(deep=True)
                    duplicates_df = all_df[all_df.duplicated(["Име*:", "Отор. код на ПОС бележка*:"], keep="first")]
                    all_df.drop_duplicates(["Име*:", "Отор. код на ПОС бележка*:"], keep="first", inplace=True)

                    whole_week = element.name.split()[1]

                    if week_to_process < 10:
                        week_to_process = "0" + str(week_to_process)

                    weekly_folder = f"Week_{week_to_process}"
                    weekly_processed_folder = f"Week_{week_to_process}/Processed files"
                    weekly_winners_folder = f"Week_{week_to_process}/Winners"

                    folders = [weekly_folder, weekly_processed_folder, weekly_winners_folder]

                    for folder in folders:
                        if not os.path.exists(folder):
                            os.makedirs(folder)

                    os.chdir(weekly_processed_folder)

                    weekly_all = f"Week_{whole_week}_all_registrations.xlsx"
                    weekly_duplicates = f"Week_{whole_week}_duplicates.xlsx"
                    weekly_without_duplicates = f"Week_{whole_week}_without_duplicates.xlsx"

                    current_files = [weekly_all, weekly_duplicates, weekly_without_duplicates]

                    for current_file in current_files:
                        if os.path.exists(current_file):
                            os.remove(current_file)

                    initial_df.to_excel(weekly_all, index=False)
                    statistics = append_and_print_statistics(f"Total registrations: {initial_df.shape[0]}", statistics)

                    duplicates_df.to_excel(weekly_duplicates, index=False)
                    statistics = append_and_print_statistics(f"Duplicates: {duplicates_df.shape[0]}", statistics)

                    all_df.to_excel(weekly_without_duplicates, index=False)
                    statistics = append_and_print_statistics(f"Without duplicates: {all_df.shape[0]}", statistics)

                    all_df["Имейл*:"] = all_df["Имейл*:"].str.strip()
                    emails_list = all_df["Имейл*:"].unique()

                    statistics = append_and_print_statistics(f"Unique emails: {len(emails_list)}\n", statistics)

                    statistics = append_and_print_statistics("Registrations by bank (without duplicates):", statistics)

                    for bank in banks_en_names:

                        bank_info_to_write = ""

                        if bank in all_df["Банка*:"].unique():
                            bank_df = all_df[all_df["Банка*:"] == bank]
                            bank_info_to_write = f"{banks_en_names[bank]}: {bank_df.shape[0]}"
                        else:
                            bank_info_to_write = f"{banks_en_names[bank]}: 0"

                        statistics = append_and_print_statistics(bank_info_to_write, statistics)

                    os.chdir("../../")
                    os.chdir(weekly_winners_folder)

                    writer = pd.ExcelWriter(f"Week_{whole_week}_winners_and_reserves.xlsx", engine="xlsxwriter")

                    weekly_drawn_df = all_df.sample(n=weekly_winners + weekly_reserves)
                    weekly_drawn_df = weekly_drawn_df.reset_index(drop=True)

                    weekly_winners_df = weekly_drawn_df.loc[0:(weekly_winners - 1)]
                    weekly_winners_df.insert(0, "Status", "Седмична награда")

                    weekly_reserves_df = weekly_drawn_df.loc[weekly_winners:(weekly_winners + weekly_reserves - 1)]
                    weekly_reserves_df.insert(0, "Status", "Седмична резерва")

                    weekly_winners_df.to_excel(writer, sheet_name="Winners", index=False)
                    weekly_reserves_df.to_excel(writer, sheet_name="Reserves", index=False)

                    writer.save()

                    winning_banks_folder = "For banks"

                    if not os.path.exists(winning_banks_folder):
                        os.makedirs(winning_banks_folder)

                    os.chdir(winning_banks_folder)

                    statistics = append_and_print_statistics("\nWinners by bank:", statistics)

                    for bank in banks_en_names:
                        bank_info_to_write = ""
                        if bank in weekly_winners_df["Банка*:"].unique():
                            winning_bank_df = weekly_winners_df[weekly_winners_df["Банка*:"] == bank]
                            bank_info_to_write = f"{banks_en_names[bank]}: {winning_bank_df.shape[0]}"
                            bank_name = banks_en_names[bank]

                            if winning_bank_df.shape[0] == 1:
                                win_bank_file = f"Week_{whole_week}_winner_{bank_name}.xlsx"
                            else:
                                win_bank_file = f"Week_{whole_week}_winners_{bank_name}.xlsx"

                            if os.path.isfile(win_bank_file):
                                os.remove(win_bank_file)

                            winning_bank_df.to_excel(win_bank_file, index=False)
                        else:
                            bank_info_to_write = f"{banks_en_names[bank]}: 0"

                        statistics = append_and_print_statistics(bank_info_to_write, statistics)

os.chdir("../../")

weekly_stats_file = f"Week_{whole_week}_statistics.txt"

if os.path.exists(weekly_stats_file):
    os.remove(weekly_stats_file)

with open(weekly_stats_file, "a+") as file:
    for row in statistics:
        file.write(f"{row}\n")

time_end = datetime.now()
time_took = time_end - time_start

print(f"\nDone! The execution of this script took {time_took.seconds} seconds.")