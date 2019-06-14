from pathlib import Path
from datetime import datetime
import pandas as pd
import re
import os
import docx
import yaml

time_start = datetime.now()

project_config = docx.Document("config.docx")

source_path = Path(project_config.paragraphs[0].text.split("=")[1].replace("\"",""))
output_path = Path(project_config.paragraphs[1].text.split("=")[1].replace("\"",""))
values_to_skip = project_config.paragraphs[2].text.split("=")[1].replace("\"","").split(",")
banks_list = yaml.safe_load(project_config.paragraphs[3].text.split("=")[1].replace("\"",""))

pattern = r"\d{2}\.\d{2}\.\d{4}"

week_to_process = int(input("Week to process: "))

all_data_frames = []
week_dates_list = []

for bank_folder in source_path.glob("*"):
    new_col_name = bank_folder.name
    for week_folder in bank_folder.iterdir():
        if int(week_folder.name.split(".")[0]) == week_to_process:
            for file in week_folder.iterdir():
                current_data_frame = pd.read_excel(file, sheet_name=0, usecols="A", header=None, skiprows=None)
                current_data_frame.rename(columns={0: "Transaction"}, inplace=True)
                current_data_frame.insert(0, "Bank", new_col_name)
                searched = re.search(pattern, file.name)
                current_data_frame.insert(2, "Date", searched[0])
                all_data_frames.append(current_data_frame)
            print(f"{new_col_name}", end=" ")
            break

all_data = pd.concat(all_data_frames)
all_data = all_data.reset_index(drop=True)

for value_to_skip in values_to_skip:
    indices_to_remove = all_data[(all_data["Transaction"] == value_to_skip)].index
    all_data.drop(indices_to_remove, inplace=True)

if week_to_process < 10:
    week_to_process = "0" + str(week_to_process)

if not os.path.exists(output_path):
    os.makedirs(output_path)

os.chdir(output_path)

if not os.path.exists(f"Week_{week_to_process}"):
    os.makedirs(f"Week_{week_to_process}")

file_with_all_data = f"Week_{week_to_process}/_All_week_{week_to_process}.csv"
all_data.to_csv(file_with_all_data, index=False)
pivot_with_all_data = f"Week_{week_to_process}/_All_week_{week_to_process}_pivot.xlsx"
#all_data["Bank"].replace(banks_list, inplace=True)
#all_data["Bank"] = all_data["Bank"].map(lambda x: x.encode("utf-8").decode("utf-8"))
pivot = pd.pivot_table(all_data, index=["Bank"], values=["Transaction"], aggfunc=len, columns=["Date"])
pivot.to_excel(pivot_with_all_data)

winners_to_draw = int(input("\nWinners per day: "))
reserves_to_draw = int(input("Reserves per day: "))
entries_to_sample = winners_to_draw + reserves_to_draw

week_dates_list = all_data["Date"].unique()

for date in week_dates_list:
    date_data_frame = all_data[all_data.Date == date]
    date_all_file = f"Week_{week_to_process}/All_{date}.csv"
    if os.path.isfile(date_all_file):
        os.remove(date_all_file)
    date_data_frame.to_csv(date_all_file, index=False, header=None)
    sample_data_frame = date_data_frame.sample(n=entries_to_sample)
    date_winners_file = f"Week_{week_to_process}/Winners_{date}.csv"
    if os.path.isfile(date_winners_file):
        os.remove(date_winners_file)
    sample_data_frame.to_csv(date_winners_file, index=False, header=None)
    print(f"{date}", end=" ")

print("\nAll files are ready!")

time_end = datetime.now()
time_took = time_end - time_start

print(f"The execution of this script took {time_took.seconds} seconds.")