import docx
import os
from pathlib import Path

cwd = os.getcwd()

project_config = docx.Document("config.docx")

source_path = Path(project_config.paragraphs[0].text.split("=")[1].replace("\"",""))
output_path = Path(project_config.paragraphs[1].text.split("=")[1].replace("\"",""))
values_to_skip = project_config.paragraphs[2].text.split("=")[1].replace("\"","").split(",")

results_folder = project_config.paragraphs[1].text.split("=")[1].replace("\"","").split("\\")[-1]


print(f"Source path: {source_path}")
print(f"Output path: {output_path}")
print(f"Values to skip: {values_to_skip}")