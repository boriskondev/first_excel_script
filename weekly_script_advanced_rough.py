import docx
import pandas as pd
from pathlib import Path
import os
import numpy as np

project_config = docx.Document("config.docx")

weekly_source_path = Path(project_config.paragraphs[1].text.split("=")[1].replace("\"", ""))
output_path = Path(project_config.paragraphs[2].text.split("=")[1].replace("\"", ""))

codes_to_remove = ['AC ', 'ac', 'АС :', 'AUTH.CODE:', 'Ac', 'АВТ. КОД', 'Авт. Код : ', 'АВТ.КОД/АС:', 'ACC', 'AUTH CODE:',
                   'АВТОМ.КОД: ', 'АВТ. КОД: ', 'Авт.код', 'авт.код.', 'autn. Code', 'AC: ', 'ac ', 'tid ', 'AUTH.CODE : ',
                   'Авт. Код:', 'Авт. Код: ', 'AUTH. CODE:', 'Auth.code:', 'AUTH. CODE', 'Авт. код: ', 'Auth.code ', 'AC:',
                   'АC', 'АС:', 'АВТ КОД ', 'AUTH. CODE: ', 'Авт.код:', 'АВТ.КОД', 'АВТ.КОД. ', 'Авт. код. ', 'АВТ КОД',
                   'АВТ.КОД: ', 'AC', 'Auth. Code ', 'AС', 'авт. Код', 'АВТ. КОД:', 'Авт. код ', 'GPA.', 'АВТ.КОД:', 'авт. код ',
                   'AC :', 'Авт.код ', 'Ас ', 'АВТ. КОД/ ', 'PRE-AUTH ', 'Авт. Код ', 'авт.код:', 'АВТ.КОД/AC:', 'Авт.код :',
                   'АС', 'Авт. код/ ']

os.chdir(output_path)

week_to_process = 1
daily_winners = 5
daily_reserves = 5
weekly_winners = 1
weekly_reserves = 5

for weekly_folder_element in weekly_source_path.glob("*"):
    if os.path.isdir(weekly_folder_element) and weekly_folder_element.name != "_Results" \
            and int(weekly_folder_element.name.split(".")[0]) == week_to_process:
        for file in weekly_folder_element.iterdir():
            if file.suffix == ".xlsx":

                all_df = pd.read_excel(file, sheet_name=0, skiprows=None)
                all_df.columns = all_df.columns.str.strip()
                for column in all_df.columns:
                    all_df[column] = all_df[column].astype('str').str.strip()

                all_df["Reg_date"], all_df["Reg_time"] = all_df["Submitted date"].str.split(" ").str

                for code in codes_to_remove:
                    all_df["Transaction Code"] = [x.strip().replace(code, '').strip() for x in all_df["Transaction Code"]]

                print(f"Total entries: {all_df.shape[0]}")

                duplicates_df = all_df[all_df.duplicated(["Firstname", "Transaction Code"], keep="first")]
                all_df.drop_duplicates(["Firstname", "Transaction Code"], keep="first", inplace=True)

                whole_week = weekly_folder_element.name.split()[1]

                weekly_duplicates = f"Week_{whole_week}_only_duplicates.xlsx"
                weekly_no_duplicates = f"Week_{whole_week}_without_duplicates.xlsx"
                weekly_fifths = f"Week_{whole_week}_fifths.xlsx"

                current_files = [weekly_duplicates, weekly_no_duplicates, weekly_fifths]

                for current_file in current_files:
                    if os.path.exists(current_file):
                        os.remove(current_file)

                duplicates_df.to_excel(weekly_duplicates, index=False)
                print(f"Duplicates: {duplicates_df.shape[0]}")
                all_df.to_excel(weekly_no_duplicates, index=False)
                print(f"Entries without duplicates: {all_df.shape[0]}")

                # Weekly prize/s logic
                fifths_df = pd.DataFrame()
                all_df["Email address (Personal)"] = all_df["Email address (Personal)"].str.strip()
                emails_list = all_df["Email address (Personal)"].unique()
                print(f"Unique emails: {len(emails_list)}")

                for email in emails_list:
                    all_selected = all_df[all_df["Email address (Personal)"] == email]
                    all_selected.index = np.arange(1, len(all_selected) + 1)
                    each_fifth = all_selected[(all_selected.index % 5 == 0)]
                    fifths_df = fifths_df.append(each_fifth)

                writer = pd.ExcelWriter(f"Week_{whole_week}_weekly_winners_and_reserves.xlsx", engine="xlsxwriter")

                weekly_drawn_df = fifths_df.sample(n=weekly_winners + weekly_reserves)
                weekly_drawn_df = weekly_drawn_df.reset_index(drop=True)
                weekly_winners_df = weekly_drawn_df.loc[0:(weekly_winners - 1)]
                weekly_winners_df.insert(0, "Status", "Седмична награда")
                weekly_reserves_df = weekly_drawn_df.loc[weekly_winners:(weekly_winners + weekly_reserves - 1)]
                weekly_reserves_df.insert(0, "Status", "Седмична резерва")

                weekly_winners_df.to_excel(writer, sheet_name="Weekly_winners", index=False)
                weekly_reserves_df.to_excel(writer, sheet_name="Weekly_reserves", index=False)

                writer.save()

                fifths_df.to_excel(weekly_fifths, index=False)

                print(f"Weekly prize entries: {fifths_df.shape[0]}")
                print(f"Weekly prize unique emails: {len(fifths_df['Email address (Personal)'].unique())}")

                # Daily prizes logic
                all_daily_winners, all_daily_reserves = [], []

                days_of_week = all_df["Reg_date"].unique()

                writer = pd.ExcelWriter(f"Week_{whole_week}_daily_winners_and_reserves.xlsx", engine="xlsxwriter")

                for day in days_of_week:
                    daily_df = all_df[all_df["Reg_date"] == day]
                    print(f"Entries for {day}: {daily_df.shape[0]}")
                    daily_drawn_df = daily_df.sample(n=daily_winners+daily_reserves)
                    daily_drawn_df = daily_drawn_df.reset_index(drop=True)
                    daily_winners_df = daily_drawn_df.loc[0:(daily_winners - 1)]
                    daily_reserves_df = daily_drawn_df.loc[daily_winners:(daily_winners + daily_reserves - 1)]
                    all_daily_winners.append(daily_winners_df)
                    all_daily_reserves.append(daily_reserves_df)

                all_daily_winners_df = pd.concat(all_daily_winners).reset_index(drop=True)
                all_daily_winners_df.insert(0, "Status", "Дневна награда")
                all_daily_reserves_df = pd.concat(all_daily_reserves).reset_index(drop=True)
                all_daily_reserves_df.insert(0, "Status", "Дневна резерва")

                all_daily_winners_df.to_excel(writer, sheet_name="Daily_winners", index=False)
                all_daily_reserves_df.to_excel(writer, sheet_name="Daily_reserves", index=False)

                writer.save()

                # All winners file
                all_winners_file = f"Week_{whole_week}_winners.xlsx"
                all_winners_df = weekly_winners_df.append(all_daily_winners_df)
                all_winners_df.to_excel(all_winners_file, index=False)

print("Done!")

# TODO
# Put in different folders
# When the script is finished - to properly name variables - what is file, what is df and so on.