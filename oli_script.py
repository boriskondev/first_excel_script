from datetime import datetime
from pathlib import Path
import docx
import pandas as pd
import os
import yaml
import numpy as np
import sys


def append_and_print_statistics(data, li):
    print(data)
    li.append(data)
    return li


def process_code(data, current_code):
    symbols = ["/", "NO"]
    data = data.strip()
    if current_code in data[0: len(current_code)]:
        data = data.strip().replace(current_code, "").strip()
        for symbol in symbols:
            if symbol in data:
                data = data.split(symbol)[0].strip()

    return data


# time_start = datetime.now()

project_config = docx.Document("config_new.docx")

SOURCE_PATH = Path(project_config.paragraphs[0].text.split("=")[1].replace("\"", ""))
OUTPUT_PATH = Path(project_config.paragraphs[1].text.split("=")[1].replace("\"", ""))
BANKS_EN_NAME = yaml.safe_load(project_config.paragraphs[2].text.split("=")[1].replace("\"", ""))
CODES_TO_REMOVE = project_config.paragraphs[3].text.split("=")[1].strip().replace("\"", "").split(", ")

statistics = []

WEEK_TO_PROCESS = 1
WEEKLY_EACH_NTH = 3
WEEKLY_WINNERS = 1
WEEKLY_RESERVES = 5

DAILY_P1 = ["Раница"]
DAILY_P1_COUNT = 3
DAILY_P2 = ["Яке"]
DAILY_P2_COUNT = 3
DAILY_P3 = ["Плажнa кърпа"]
DAILY_P3_COUNT = 7

DAILY_PRIZES_LIST = DAILY_P1 * DAILY_P1_COUNT + DAILY_P2 * DAILY_P2_COUNT + DAILY_P3 * DAILY_P3_COUNT

DAILY_WINNERS = DAILY_P1_COUNT + DAILY_P2_COUNT + DAILY_P3_COUNT
DAILY_RESERVES = DAILY_WINNERS

whole_week = ""

os.chdir(OUTPUT_PATH)

for element in SOURCE_PATH.glob("*"):
    if os.path.isdir(element):
        if element.name != "_Results" and element.name != "Additional" \
                and int(element.name.split(".")[0]) == WEEK_TO_PROCESS:
            for file in element.iterdir():
                if file.suffix == ".xlsx":
                    all_df = pd.read_excel(file, sheet_name=0, skiprows=None)
                    all_df.columns = all_df.columns.str.strip()
                    all_df["Email address (Personal)"] = all_df["Email address (Personal)"].str.strip()
                    all_df["First name"] = all_df["First name"].str.strip()
                    all_df["Bank"] = all_df["Bank"].str.strip()
                    all_df[["Reg_date", "Reg_time"]] = all_df["Submitted date"].str.split(" ", expand=True)

                    for code in CODES_TO_REMOVE:
                        all_df["Receipt number"] = \
                            [process_code(x, code) for x in all_df["Receipt number"]]

                    initial_df = all_df.copy(deep=True)
                    # Check if date to be added as third criteria
                    duplicates_df = all_df[all_df.duplicated(["First name", "Receipt number"], keep="first")]
                    all_df.drop_duplicates(["First name", "Receipt number"], keep="first", inplace=True)

                    whole_week = element.name.split()[1]

                    if WEEK_TO_PROCESS < 10:
                        WEEK_TO_PROCESS = "0" + str(WEEK_TO_PROCESS)

                    weekly_folder = f"Week_{WEEK_TO_PROCESS}"
                    weekly_processed_folder = f"Week_{WEEK_TO_PROCESS}/Processed files"
                    weekly_winners_folder = f"Week_{WEEK_TO_PROCESS}/Winners"

                    folders = [weekly_folder, weekly_processed_folder, weekly_winners_folder]

                    for folder in folders:
                        if not os.path.exists(folder):
                            os.makedirs(folder)

                    os.chdir(weekly_processed_folder)

                    weekly_all = f"Week_{whole_week}_all_registrations.xlsx"
                    weekly_duplicates = f"Week_{whole_week}_duplicates.xlsx"
                    weekly_without_duplicates = f"Week_{whole_week}_without_duplicates.xlsx"
                    weekly_thirds = f"Week_{whole_week}_thirds.xlsx"

                    current_files = [weekly_all, weekly_duplicates, weekly_without_duplicates, weekly_thirds]

                    for current_file in current_files:
                        if os.path.exists(current_file):
                            os.remove(current_file)

                    initial_df.to_excel(weekly_all, index=False)
                    # statistics = append_and_print_statistics(f"Total registrations: {initial_df.shape[0]}", statistics)

                    duplicates_df.to_excel(weekly_duplicates, index=False)
                    # statistics = append_and_print_statistics(f"Duplicates: {duplicates_df.shape[0]}", statistics)

                    all_df.to_excel(weekly_without_duplicates, index=False)
                    # statistics = append_and_print_statistics(f"Without duplicates: {all_df.shape[0]}", statistics)

                    # Weekly prize logic
                    thirds_df = pd.DataFrame()
                    emails_list = all_df["Email address (Personal)"].unique()
                    # statistics = append_and_print(f"Unique emails: {len(emails_list)}", statistics)

                    for email in emails_list:
                        all_selected = all_df[all_df["Email address (Personal)"] == email]
                        all_selected.index = np.arange(1, len(all_selected) + 1)
                        each_third = all_selected[(all_selected.index % WEEKLY_EACH_NTH == 0)]
                        thirds_df = thirds_df.append(each_third)

                    thirds_df.to_excel(weekly_thirds, index=False)

                    os.chdir("../../")
                    os.chdir(weekly_winners_folder)

                    writer = pd.ExcelWriter(f"Week_{whole_week}_weekly_winners_and_reserves.xlsx", engine="xlsxwriter")

                    weekly_drawn_df = thirds_df.sample(n=WEEKLY_WINNERS + WEEKLY_RESERVES)
                    weekly_drawn_df = weekly_drawn_df.reset_index(drop=True)
                    weekly_winners_df = weekly_drawn_df.loc[0:(WEEKLY_WINNERS - 1)]
                    weekly_winners_df.insert(0, "Status", "Седмична награда")
                    weekly_reserves_df = weekly_drawn_df.loc[WEEKLY_WINNERS:(WEEKLY_WINNERS + WEEKLY_RESERVES - 1)]
                    weekly_reserves_df.insert(0, "Status", "Седмична резерва")

                    weekly_winners_df.to_excel(writer, sheet_name="Weekly_winners", index=False)
                    weekly_reserves_df.to_excel(writer, sheet_name="Weekly_reserves", index=False)

                    writer.save()

                    # statistics = append_and_print(f"Weekly prize entries: {fifths_df.shape[0]}", statistics)
                    # unique_emails_count = len(thirds_df['Email address (Personal)'].unique())
                    # statistics = append_and_print(f"Weekly prize unique emails: {unique_emails_count}", statistics)

                    # Daily prizes logic
                    all_daily_winners, all_daily_reserves = [], []

                    days_of_week = all_df["Reg_date"].unique()

                    writer = pd.ExcelWriter(f"Week_{whole_week}_daily_winners_and_reserves.xlsx", engine="xlsxwriter")

                    for day in days_of_week:
                        daily_df = all_df[all_df["Reg_date"] == day]
                        # statistics = append_and_print(f"Entries for {day}: {daily_df.shape[0]}", statistics)
                        daily_drawn_df = daily_df.sample(n=DAILY_WINNERS + DAILY_RESERVES)
                        daily_drawn_df = daily_drawn_df.reset_index(drop=True)
                        daily_winners_df = daily_drawn_df.loc[0:(DAILY_WINNERS - 1)]
                        daily_winners_df.insert(0, "Status", DAILY_PRIZES_LIST)
                        daily_reserves_df = daily_drawn_df.loc[DAILY_WINNERS:(DAILY_WINNERS + DAILY_RESERVES - 1)]
                        all_daily_winners.append(daily_winners_df)
                        all_daily_reserves.append(daily_reserves_df)

                    all_daily_winners_df = pd.concat(all_daily_winners).reset_index(drop=True)
                    all_daily_reserves_df = pd.concat(all_daily_reserves).reset_index(drop=True)
                    all_daily_reserves_df.insert(0, "Status", "Дневна резерва")

                    all_daily_winners_df.to_excel(writer, sheet_name="Daily_winners", index=False)
                    all_daily_reserves_df.to_excel(writer, sheet_name="Daily_reserves", index=False)

                    writer.save()


                    # winning_banks_folder = "For banks"
                    #
                    # if not os.path.exists(winning_banks_folder):
                    #     os.makedirs(winning_banks_folder)
                    #
                    # os.chdir(winning_banks_folder)
                    #
                    # statistics = append_and_print_statistics("\nWinners by bank:", statistics)
                    #
                    # for bank in banks_en_names:
                    #     bank_info_to_write = ""
                    #     if bank in weekly_winners_df["Банка*:"].unique():
                    #         winning_bank_df = weekly_winners_df[weekly_winners_df["Банка*:"] == bank]
                    #         bank_info_to_write = f"{banks_en_names[bank]}: {winning_bank_df.shape[0]}"
                    #         bank_name = banks_en_names[bank]
                    #
                    #         if winning_bank_df.shape[0] == 1:
                    #             win_bank_file = f"Week_{whole_week}_winner_{bank_name}.xlsx"
                    #         else:
                    #             win_bank_file = f"Week_{whole_week}_winners_{bank_name}.xlsx"
                    #
                    #         if os.path.isfile(win_bank_file):
                    #             os.remove(win_bank_file)
                    #
                    #         winning_bank_df.to_excel(win_bank_file, index=False)
                    #     else:
                    #         bank_info_to_write = f"{banks_en_names[bank]}: 0"
                    #
                    #     statistics = append_and_print_statistics(bank_info_to_write, statistics)

# os.chdir("../../")
#
# weekly_stats_file = f"Week_{whole_week}_statistics.txt"
#
# if os.path.exists(weekly_stats_file):
#     os.remove(weekly_stats_file)
#
# with open(weekly_stats_file, "a+") as file:
#     for row in statistics:
#         file.write(f"{row}\n")
#
# time_end = datetime.now()
# time_took = time_end - time_start
#
# print(f"\nDone! The execution of this script took {time_took.seconds} seconds.")